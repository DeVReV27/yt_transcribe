import streamlit as st
from openai import OpenAI
import os
import yt_dlp
import docx
import tempfile
import json
import time
import shutil
import io
from pathlib import Path
from pydub import AudioSegment
import math
import re

# Get API key from Streamlit secrets
api_key = st.secrets["OPENAI_API_KEY"]

# Initialize OpenAI client
client = OpenAI(api_key=api_key)

# Initialize session state
if 'chat_history' not in st.session_state:
    st.session_state.chat_history = []
if 'current_transcript' not in st.session_state:
    st.session_state.current_transcript = ""
if 'generated_content' not in st.session_state:
    st.session_state.generated_content = ""
if 'video_title' not in st.session_state:
    st.session_state.video_title = ""
if 'channel_name' not in st.session_state:
    st.session_state.channel_name = ""
if 'transcribed_text' not in st.session_state:
    st.session_state.transcribed_text = ""

def get_video_info(url):
    """Get video title and channel name from YouTube URL"""
    try:
        ydl_opts = {
            'quiet': True,
            'no_warnings': True,
            'extract_flat': True
        }
        with yt_dlp.YoutubeDL(ydl_opts) as ydl:
            info = ydl.extract_info(url, download=False)
            return info.get('title', ''), info.get('channel', '')
    except Exception as e:
        st.error(f"Error getting video info: {str(e)}")
        return '', ''

def clean_temp_files(file_path):
    """Clean up temporary files"""
    try:
        if file_path and os.path.exists(file_path):
            os.remove(file_path)
    except Exception as e:
        st.warning(f"Warning: Could not clean up temporary file: {str(e)}")

def split_audio(audio_path, chunk_duration=300000):
    """Split audio file into smaller chunks"""
    try:
        audio = AudioSegment.from_mp3(audio_path)
        chunks = []
        
        # Calculate number of chunks needed
        total_duration = len(audio)
        num_chunks = math.ceil(total_duration / chunk_duration)
        
        for i in range(num_chunks):
            start_time = i * chunk_duration
            end_time = min((i + 1) * chunk_duration, total_duration)
            
            # Extract chunk
            chunk = audio[start_time:end_time]
            
            # Save chunk to temporary file
            chunk_path = f"{audio_path}_chunk_{i}.mp3"
            chunk.export(chunk_path, format="mp3")
            chunks.append(chunk_path)
        
        return chunks
    except Exception as e:
        st.error(f"Error splitting audio: {str(e)}")
        return None

def download_youtube_audio(url):
    """Download audio from YouTube video using yt-dlp"""
    output_path = None
    try:
        # Validate YouTube URL
        if not re.match(r'^(https?://)?(www\.)?(youtube\.com|youtu\.be)/.+$', url):
            st.error("Please enter a valid YouTube URL")
            return None

        # Create a unique temporary directory
        temp_dir = tempfile.mkdtemp()
        output_path = os.path.join(temp_dir, "audio")
        
        ydl_opts = {
            'format': 'bestaudio/best',
            'outtmpl': output_path + '.%(ext)s',
            'postprocessors': [{
                'key': 'FFmpegExtractAudio',
                'preferredcodec': 'mp3',
                'preferredquality': '192',
            }],
            'quiet': True,
            'no_warnings': True
        }
        
        with yt_dlp.YoutubeDL(ydl_opts) as ydl:
            ydl.download([url])
            
        return output_path + '.mp3'
    except Exception as e:
        if output_path and os.path.exists(output_path):
            clean_temp_files(output_path + '.mp3')
        st.error(f"Error downloading YouTube video: {str(e)}")
        st.info("If the error persists, try another video or check if the video is publicly accessible")
        return None

def transcribe_chunk(chunk_path, max_retries=3):
    """Transcribe a single audio chunk"""
    for attempt in range(max_retries):
        try:
            with open(chunk_path, "rb") as file:
                transcript = client.audio.transcriptions.create(
                    model="whisper-1",
                    file=file,
                    response_format="text"
                )
            return transcript
        except Exception as e:
            if attempt == max_retries - 1:
                st.error(f"Error during transcription: {str(e)}")
                return None
            time.sleep(2)  # Wait before retrying
    return None

def transcribe_audio(audio_file, max_retries=3):
    """Transcribe audio file using OpenAI Whisper with chunking"""
    try:
        # Split audio into chunks
        chunks = split_audio(audio_file)
        if not chunks:
            return None

        # Initialize progress bar
        progress_bar = st.progress(0)
        status_text = st.empty()
        
        # Transcribe each chunk
        transcripts = []
        for i, chunk_path in enumerate(chunks):
            status_text.text(f"Transcribing part {i+1} of {len(chunks)}...")
            transcript = transcribe_chunk(chunk_path, max_retries)
            if transcript:
                transcripts.append(transcript)
            progress_bar.progress((i + 1) / len(chunks))
            clean_temp_files(chunk_path)
        
        # Combine all transcripts
        final_transcript = " ".join(transcripts)
        status_text.empty()
        progress_bar.empty()
        
        return final_transcript
                
    except Exception as e:
        st.error(f"Error transcribing audio: {str(e)}")
        return None
    finally:
        # Clean up the temporary audio file
        clean_temp_files(audio_file)

def chunk_text(text, chunk_size=2000):
    """Split text into chunks of approximately equal size"""
    words = text.split()
    chunks = []
    current_chunk = []
    current_size = 0
    
    for word in words:
        current_chunk.append(word)
        current_size += 1
        if current_size >= chunk_size:
            chunks.append(' '.join(current_chunk))
            current_chunk = []
            current_size = 0
    
    if current_chunk:
        chunks.append(' '.join(current_chunk))
    
    return chunks

def generate_new_content(text, temperature=0.7):
    """Generate new content using GPT-4 with text chunking"""
    try:
        # Split text into chunks
        chunks = chunk_text(text)
        generated_chunks = []
        
        # Process each chunk
        for i, chunk in enumerate(chunks):
            # Calculate word count for this chunk
            chunk_word_count = len(chunk.split())
            
            response = client.chat.completions.create(
                model="gpt-4",
                messages=[
                    {"role": "system", "content": f"""You are a creative writer. Rewrite the given text in a unique way while:
                    1. Maintaining the core message and key points
                    2. Avoiding plagiarism
                    3. Keeping approximately the same length (target: {chunk_word_count} words)
                    4. Preserving the depth and detail of the original content
                    5. Using a similar structure but with fresh language
                    6. Ensuring the text flows smoothly with other chunks"""},
                    {"role": "user", "content": f"Please rewrite this text (part {i+1} of {len(chunks)}): {chunk}"}
                ],
                temperature=temperature,
                max_tokens=3000
            )
            generated_chunks.append(response.choices[0].message.content)
        
        # Combine all chunks
        return ' '.join(generated_chunks)
    except Exception as e:
        st.error(f"Error generating content: {str(e)}")
        return None

def chat_with_gpt(prompt, history, temperature=0.7):
    """Chat with GPT-4 for content refinement"""
    try:
        messages = [
            {"role": "system", "content": "You are a helpful content editor, helping users refine and modify their text based on their requirements."}
        ]
        
        # Add chat history
        for msg in history:
            messages.append({"role": "user" if msg["is_user"] else "assistant", "content": msg["message"]})
        
        # Add new prompt
        messages.append({"role": "user", "content": prompt})
        
        response = client.chat.completions.create(
            model="gpt-4",
            messages=messages,
            temperature=temperature
        )
        return response.choices[0].message.content
    except Exception as e:
        st.error(f"Error in chat: {str(e)}")
        return None

# Streamlit UI
st.title("YouTube Video Transcription & Content Generator")

# Sidebar
with st.sidebar:
    st.title("Options")
    
    # Display video title and channel name if available
    if st.session_state.video_title and st.session_state.channel_name:
        st.markdown(f"**Video:** {st.session_state.video_title}")
        st.markdown(f"**Channel:** {st.session_state.channel_name}")
    
    # Mode selection
    mode = st.radio("Select Mode", ["Transcribe Video", "Create New Content"])
    
    # Temperature slider
    temperature = st.slider("AI Temperature", 0.0, 1.0, 0.7)
    
    # File upload for text rewriting
    uploaded_file = st.file_uploader("Upload text file for rewriting", type=['txt', 'docx'])
    
    # Add some spacing
    st.markdown("<br>", unsafe_allow_html=True)
    
    # Add a footer
    st.markdown("---")
    st.markdown("© 2024 MacieraiTranscribe. All rights reserved.")

# Main content area
if mode == "Transcribe Video":
    st.header("Video Transcription")
    video_url = st.text_input("Enter YouTube Video URL")
    
    if video_url:
        if st.button("Transcribe"):
            with st.spinner("Processing video..."):
                # Get video information
                title, channel = get_video_info(video_url)
                if title and channel:
                    st.session_state.video_title = title
                    st.session_state.channel_name = channel
                
                st.info("Downloading video... This may take a moment.")
                audio_file = download_youtube_audio(video_url)
                if audio_file:
                    st.info("Download successful! Now transcribing...")
                    transcript = transcribe_audio(audio_file)
                    if transcript:
                        st.session_state.transcribed_text = transcript
                        st.success("Transcription complete!")
                        st.text_area("Transcript", transcript, height=300)
                        
                        # Add download buttons directly under the transcript
                        col1, col2 = st.columns(2)
                        with col1:
                            if st.download_button(
                                "Download as Text",
                                transcript,
                                file_name="transcript.txt",
                                mime="text/plain"
                            ):
                                st.success("Text file downloaded!")
                        with col2:
                            # Create Word document
                            doc = docx.Document()
                            doc.add_paragraph(transcript)
                            bio = io.BytesIO()
                            doc.save(bio)
                            if st.download_button(
                                label="Download as Word",
                                data=bio.getvalue(),
                                file_name="transcript.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            ):
                                st.success("Word document downloaded!")

elif mode == "Create New Content":
    st.header("Content Generation")
    
    # Handle uploaded file
    if uploaded_file:
        if uploaded_file.type == "text/plain":
            content = uploaded_file.read().decode()
        else:  # docx
            doc = docx.Document(uploaded_file)
            content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        st.session_state.current_transcript = content
    
    # Display current transcript
    if st.session_state.current_transcript:
        st.subheader("Original Text")
        st.text_area("Original", st.session_state.current_transcript, height=200)
        
        if st.button("Generate New Content"):
            with st.spinner("Generating new content..."):
                new_content = generate_new_content(st.session_state.current_transcript, temperature)
                if new_content:
                    st.session_state.generated_content = new_content
                    st.success("Content generated!")
    
    # Display generated content
    if st.session_state.generated_content:
        st.subheader("Generated Content")
        st.text_area("Generated", st.session_state.generated_content, height=300)
        
        # Add word count comparison
        original_words = len(st.session_state.current_transcript.split())
        generated_words = len(st.session_state.generated_content.split())
        st.info(f"Word count - Original: {original_words}, Generated: {generated_words}")
        
        # Download buttons for generated content
        col1, col2 = st.columns(2)
        with col1:
            if st.download_button(
                "Download as Text",
                st.session_state.generated_content,
                file_name="generated_content.txt",
                mime="text/plain"
            ):
                st.success("Text file downloaded!")
        with col2:
            doc = docx.Document()
            doc.add_paragraph(st.session_state.generated_content)
            bio = io.BytesIO()
            doc.save(bio)
            if st.download_button(
                label="Download as Word",
                data=bio.getvalue(),
                file_name="generated_content.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            ):
                st.success("Word document downloaded!")
        
        # Chat interface for content refinement
        st.subheader("Content Refinement Chat")
        user_input = st.text_input("Enter your refinement request (e.g., 'make it more formal' or 'make it shorter')")
        
        if user_input:
            if st.button("Send"):
                with st.spinner("Processing..."):
                    response = chat_with_gpt(user_input, st.session_state.chat_history, temperature)
                    if response:
                        st.session_state.chat_history.append({"is_user": True, "message": user_input})
                        st.session_state.chat_history.append({"is_user": False, "message": response})
                        st.session_state.generated_content = response
        
        # Display chat history
        for message in st.session_state.chat_history:
            if message["is_user"]:
                st.write("You:", message["message"])
            else:
                st.write("Assistant:", message["message"])

# Add footer with instructions
st.markdown("---")
st.markdown("""
### Instructions:
1. **Transcribe Video Mode**:
   - Paste a YouTube video URL and click 'Transcribe'
   - The application will automatically handle videos of any length by processing them in chunks
   - Progress will be shown during transcription
   - Download the transcript when complete

2. **Create New Content Mode**:
   - Either use a transcribed video or upload a text/Word file
   - Adjust the AI temperature (higher = more creative)
   - Generate new content and refine it using the chat interface
   - Download the final content when ready
""")

# Clean up any remaining temporary files when the app is closed
if hasattr(st.session_state, 'temp_files'):
    for file in st.session_state.temp_files:
        clean_temp_files(file)
